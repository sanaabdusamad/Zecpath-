from docx import Document

def read_docx(file_path: str) -> str:
    doc = Document(file_path)
    lines = []

    # paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            lines.append(p.text)

    # tables (optional but useful)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                lines.append(row_text)

    return "\n".join(lines)